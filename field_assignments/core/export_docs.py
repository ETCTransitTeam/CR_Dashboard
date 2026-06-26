from __future__ import annotations

import re
import zipfile
from collections import OrderedDict
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook

from field_assignments.core.constants import COLUMN_WIDTHS
from field_assignments.core.time_utils import (
    display_value,
    format_time,
    normalize_assignment,
    normalize_header,
    subtract_minutes,
    time_to_minutes,
)
from field_assignments.core.workbook import validate_headers


def unique_values(rows: Iterable[dict[str, object]], key: str) -> list[str]:
    seen: list[str] = []
    for row in rows:
        value = display_value(key, row.get(key))
        if value and value not in seen:
            seen.append(value)
    return seen


def safe_filename(text: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*#]+', "_", text)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("._")
    return cleaned or "assignment"


def compact_time_for_filename(value: object) -> str:
    from field_assignments.core.time_utils import parse_time

    parsed = parse_time(value)
    if parsed is None:
        return safe_filename(str(value or "")).lower().replace("_", "")
    return parsed.strftime("%I%M%p").lower()


def assignment_filename(assignment: str, rows: list[dict[str, object]]) -> str:
    rows = sorted(
        rows,
        key=lambda row: (
            time_to_minutes(row.get("Start Time")) is None,
            time_to_minutes(row.get("Start Time")) or 0,
        ),
    )
    first_row = rows[0]
    last_row = rows[-1]
    block = "_".join(safe_filename(value) for value in unique_values(rows, "Block"))
    report_location = safe_filename(display_value("Start Location", first_row.get("Start Location")))
    shift = (
        f"{compact_time_for_filename(first_row.get('Start Time'))}-"
        f"{compact_time_for_filename(last_row.get('End Time'))}"
    )
    return f"{block}_Assignment-{safe_filename(assignment)}_{report_location}_{shift}.docx"


def read_assignments(source: bytes | Path, sheet_name: str | None = None) -> tuple[list[str], OrderedDict[str, list[dict[str, object]]]]:
    if isinstance(source, Path):
        workbook = load_workbook(source, data_only=True)
    else:
        workbook = load_workbook(BytesIO(source), data_only=True)
    sheet = workbook[sheet_name] if sheet_name else workbook.active

    headers = [normalize_header(sheet.cell(row=1, column=i).value) for i in range(1, 9)]
    validate_headers(headers)

    assignments: OrderedDict[str, list[dict[str, object]]] = OrderedDict()
    for excel_row in sheet.iter_rows(min_row=2, max_col=8, values_only=True):
        if not any(value is not None and str(value).strip() != "" for value in excel_row):
            continue
        assignment = normalize_assignment(excel_row[0])
        if assignment is None:
            continue
        row = dict(zip(headers, excel_row))
        assignments.setdefault(assignment, []).append(row)

    return headers, assignments


def assignment_sort_key(item: tuple[str, list[dict[str, object]]]) -> tuple[int, str]:
    assignment, _rows = item
    try:
        return (0, f"{int(assignment):06d}")
    except ValueError:
        return (1, assignment)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "8A8A8A", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_labeled_header_line(container, label: str, value: str) -> None:
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    label_run = paragraph.add_run(f"{label} ")
    label_run.font.name = "Arial"
    label_run.font.size = Pt(10)
    value_run = paragraph.add_run(value)
    value_run.bold = True
    value_run.font.name = "Arial"
    value_run.font.size = Pt(10)


def build_document(assignment: str, rows: list[dict[str, object]], headers: list[str], output_path: Path | None = None) -> bytes:
    rows = sorted(
        rows,
        key=lambda row: (
            time_to_minutes(row.get("Start Time")) is None,
            time_to_minutes(row.get("Start Time")) or 0,
        ),
    )
    first_row = rows[0]
    last_row = rows[-1]

    report_location = display_value("Start Location", first_row.get("Start Location"))
    first_start = first_row.get("Start Time")
    last_end = last_row.get("End Time")
    shift = f"{format_time(first_start)} - {format_time(last_end)}"
    blocks = ", ".join(unique_values(rows, "Block"))

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    header = section.header
    if header.paragraphs:
        header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)
    add_labeled_header_line(header, "Assignment #:", assignment)
    add_labeled_header_line(header, "Report Location:", report_location)
    add_labeled_header_line(header, "Report Time:", subtract_minutes(first_start, 15))
    add_labeled_header_line(header, "Block:", blocks)
    add_labeled_header_line(header, "Shift:", shift)

    footer = section.footer
    if footer.paragraphs:
        footer.paragraphs[0].text = ""
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(120)

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, header_text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(COLUMN_WIDTHS[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F2F2")
        set_cell_margins(cell, top=90, bottom=90, start=70, end=70)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.name = "Arial Narrow"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)

    for row_data in rows:
        cells = table.add_row().cells
        for index, header_text in enumerate(headers):
            cell = cells[index]
            cell.width = Inches(COLUMN_WIDTHS[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=70, bottom=70, start=70, end=70)
            paragraph = cell.paragraphs[0]
            if header_text in {"Asn#", "Block", "Route", "Start Time", "End Time"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(display_value(header_text, row_data.get(header_text)))
            run.font.name = "Arial Narrow"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)

    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(data)
    return data


def export_reports_zip(
    source: bytes,
    sheet_name: str | None = None,
    wanted: set[str] | None = None,
) -> tuple[bytes, list[dict[str, str]]]:
    headers, assignments = read_assignments(source, sheet_name)
    summaries: list[dict[str, str]] = []
    archive = BytesIO()

    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for assignment, rows in sorted(assignments.items(), key=assignment_sort_key):
            if wanted is not None and assignment not in wanted:
                continue
            sorted_rows = sorted(
                rows,
                key=lambda row: (
                    time_to_minutes(row.get("Start Time")) is None,
                    time_to_minutes(row.get("Start Time")) or 0,
                ),
            )
            filename = assignment_filename(assignment, sorted_rows)
            doc_bytes = build_document(assignment, sorted_rows, headers)
            zf.writestr(filename, doc_bytes)
            first_row = sorted_rows[0]
            last_row = sorted_rows[-1]
            summaries.append(
                {
                    "Assignment": assignment,
                    "Block": ", ".join(unique_values(sorted_rows, "Block")),
                    "Report Location": display_value("Start Location", first_row.get("Start Location")),
                    "Report Time": subtract_minutes(first_row.get("Start Time"), 15),
                    "Shift": f"{format_time(first_row.get('Start Time'))} - {format_time(last_row.get('End Time'))}",
                    "Rows": str(len(sorted_rows)),
                    "File Name": filename,
                }
            )

    if not summaries:
        raise ValueError("No assignments matched the export criteria.")

    return archive.getvalue(), summaries
